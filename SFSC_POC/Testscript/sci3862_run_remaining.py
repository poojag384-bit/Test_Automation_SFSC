#!/usr/bin/env python3
"""
SCI-3862 — run steps 16–63 in QA (sixt3--qa): BNL_VIP, Overflow_VIP,
FR_T3_During_Rental_Loyalty, US_T3_After_Rental_Loyalty.

Prereqs:
  pip install playwright python-docx
  python -m playwright install chromium

Usage:
  python3 sci3862_run_remaining.py

A browser window opens. Log into Salesforce QA if prompted, then return to the
terminal and press Enter so the script can continue from the case page.

Adjust QUEUE_SKILLS if a backlog row does not appear (pick skills that match
the queue in Command Center > Skills filter list).
"""

from __future__ import annotations

import argparse
import re
import sys
import time
from pathlib import Path

from docx import Document
from docx.shared import Inches

try:
    from playwright.sync_api import sync_playwright
except ImportError:
    print("Install: pip install playwright && python -m playwright install chromium", file=sys.stderr)
    raise

# --- constants ---
CASE_URL = (
    "https://sixt3--qa.sandbox.lightning.force.com/lightning/r/Case/500Fg00000mqxy8IAA/view"
)
COMMAND_CENTER_HASH = (
    "https://sixt3--qa.sandbox.lightning.force.com/one/one.app"
    "#eyJjb21wb25lbnREZWYiOiJvbW5pOnN1cGVydmlzb3JQYW5lbCIsImF0dHJpYnV0ZXMiOnt9LCJzdGF0ZSI6e319"
)
CASE_NO = "15856096"
OUT = Path(__file__).resolve().parent

# Queue display name -> skills to check in Skills Backlog filter (must exist in org)
QUEUE_SKILLS: list[tuple[str, list[str], str]] = [
    ("BNL_VIP", ["Netherlands", "English", "VIP"], "BNL_VIP"),
    ("Overflow_VIP", ["France", "English", "VIP"], "Overflow_VIP"),
    (
        "FR_T3_During_Rental_Loyalty",
        ["France", "English", "T3_VIP"],
        "FR_T3_During_Rental_Loyalty",
    ),
    (
        "US_T3_After_Rental_Loyalty",
        ["United States", "English", "T3_VIP"],
        "US_T3_After_Rental_Loyalty",
    ),
]


def _shot(page, name: str) -> Path:
    p = OUT / f"{name}.png"
    page.screenshot(path=str(p))
    return p


def change_case_owner(page, queue_name: str) -> None:
    page.get_by_role("button", name="Change Owner").click()
    page.wait_for_timeout(2000)
    combo = page.get_by_role("combobox", name="Select Owner")
    combo.click()
    combo.fill(queue_name)
    page.wait_for_timeout(1500)
    adv = page.get_by_role("dialog", name="Advanced Search")
    adv.locator("label.slds-radio__label").first.click()
    page.wait_for_timeout(400)
    adv.get_by_role("button", name="Select", exact=True).click()
    page.wait_for_timeout(1000)
    dlg = page.get_by_role("dialog").filter(has=page.get_by_role("heading", name="Change Owner"))
    dlg.get_by_role("button", name="Change Owner").click()
    page.wait_for_timeout(4000)


def set_case_priority(page, label: str) -> None:
    page.get_by_role("button", name="Edit Priority").first.click()
    page.wait_for_timeout(1200)
    combo = page.get_by_role("combobox", name=re.compile("priority", re.I))
    combo.click()
    page.wait_for_timeout(400)
    page.get_by_role("option", name=re.compile(f"^{re.escape(label)}$", re.I)).click()
    page.wait_for_timeout(400)
    page.get_by_role("button", name=re.compile("^Save$", re.I)).click()
    page.wait_for_timeout(3000)


def open_skills_backlog_filtered(page, skills: list[str]) -> None:
    page.goto(COMMAND_CENTER_HASH)
    page.wait_for_timeout(2500)
    page.get_by_role("link", name="Command Center for Service").click()
    page.wait_for_timeout(1500)
    page.get_by_role("tab", name="Skills Backlog").click()
    page.wait_for_timeout(2000)
    try:
        page.get_by_role("button", name="Clear Filters").click(timeout=3000)
    except Exception:
        pass
    page.wait_for_timeout(500)
    page.get_by_role("button", name="Filter for column SKILLS").click()
    page.wait_for_timeout(800)
    dlg = page.get_by_role("dialog").filter(has_text="Show work items that require")
    dlg.get_by_role("checkbox", name="All").uncheck()
    page.wait_for_timeout(400)
    for s in skills:
        dlg.get_by_role("checkbox", name=s, exact=True).check()
    page.get_by_role("heading", name="Skills Backlog Summary").click()
    page.wait_for_timeout(5000)


def run_phase(page, queue: str, skills: list[str], doc_base: str) -> list[Path]:
    """Steps mirror script: owner+screenshot, medium+backlog55, low+backlog57, high case+backlog51."""
    paths: list[Path] = []
    page.goto(CASE_URL)
    page.wait_for_timeout(3000)

    change_case_owner(page, queue)
    paths.append(_shot(page, f"sci3862_step_owner_{doc_base}"))

    set_case_priority(page, "Medium")
    open_skills_backlog_filtered(page, skills)
    paths.append(_shot(page, f"sci3862_{doc_base}_backlog_medium_55"))

    page.goto(CASE_URL)
    page.wait_for_timeout(2000)
    set_case_priority(page, "Low")
    open_skills_backlog_filtered(page, skills)
    paths.append(_shot(page, f"sci3862_{doc_base}_backlog_low_57"))

    page.goto(CASE_URL)
    page.wait_for_timeout(2000)
    set_case_priority(page, "High")
    paths.append(_shot(page, f"sci3862_{doc_base}_case_high"))

    open_skills_backlog_filtered(page, skills)
    paths.append(_shot(page, f"sci3862_{doc_base}_backlog_high_51"))

    page.goto(CASE_URL)
    page.wait_for_timeout(2000)
    set_case_priority(page, "Medium")
    return paths


def build_docx(title: str, doc_name: str, images: list[Path]) -> Path:
    doc = Document()
    doc.add_heading(f"SCI-3862 — {title}", 0)
    doc.add_paragraph(f"Case {CASE_NO} | QA Sandbox")
    for p in images:
        if p.is_file():
            doc.add_heading(p.stem, level=1)
            doc.add_picture(str(p), width=Inches(6.2))
    out = OUT / f"{doc_name}.docx"
    doc.save(out)
    return out


def main() -> None:
    ap = argparse.ArgumentParser(description="SCI-3862 steps 16–63")
    ap.add_argument(
        "--auto",
        action="store_true",
        help="Skip 'press Enter' (waits 8s for you to focus the browser / log in)",
    )
    ap.add_argument("--headless", action="store_true", help="Run without browser UI")
    args = ap.parse_args()

    headless = args.headless
    print("Opening browser. Log in to Salesforce QA if needed.")
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=headless)
        context = browser.new_context(viewport={"width": 1440, "height": 900})
        page = context.new_page()
        page.goto(CASE_URL, wait_until="domcontentloaded", timeout=120000)
        if args.auto:
            print("Auto mode: continuing in 8 seconds…")
            time.sleep(8)
        else:
            input("After you can see the case (logged in), press Enter here to start… ")

        all_docs: list[tuple[str, str, list[Path]]] = []

        for queue, skills, doc_base in QUEUE_SKILLS:
            print(f"=== Phase: {queue} ===")
            imgs = run_phase(page, queue, skills, doc_base)
            all_docs.append((queue.replace("_", " "), doc_base, imgs))

        for title, doc_base, imgs in all_docs:
            out = build_docx(title, doc_base, imgs)
            print(f"Wrote {out}")

        browser.close()
    print("Done.")


if __name__ == "__main__":
    main()
